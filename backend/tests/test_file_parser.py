from pathlib import Path

from docx import Document
import pytest
from fastapi import HTTPException

from novel_deconstructor.services import file_parser
from novel_deconstructor.services.file_parser import normalize_text_file, validate_extension


def test_txt_parser_utf8(tmp_path: Path):
    source = tmp_path / "novel.txt"
    target = tmp_path / "out.txt"
    source.write_text("第一章\n你好\n", encoding="utf-8")

    meta = normalize_text_file(source, target)

    assert meta["encoding"] in {"utf-8", "utf-8-sig"}
    assert target.read_text(encoding="utf-8") == "第一章\n你好\n"


def test_txt_parser_gbk(tmp_path: Path):
    source = tmp_path / "novel.txt"
    target = tmp_path / "out.txt"
    source.write_bytes("第一章\n你好\n".encode("gbk"))

    meta = normalize_text_file(source, target)

    assert meta["encoding"] in {"gb18030", "gbk"}
    assert "第一章" in target.read_text(encoding="utf-8")


def test_markdown_extension_supported():
    assert validate_extension("draft.md") == "md"


def test_docx_parser(tmp_path: Path):
    source = tmp_path / "novel.docx"
    target = tmp_path / "out.txt"
    doc = Document()
    doc.add_paragraph("第一章 风起")
    doc.add_paragraph("你好，世界。")
    doc.save(source)

    meta = normalize_text_file(source, target)

    assert meta["parser"] == "docx"
    text = target.read_text(encoding="utf-8")
    assert "第一章 风起" in text
    assert "你好，世界。" in text


def test_pdf_parser(monkeypatch: pytest.MonkeyPatch, tmp_path: Path):
    source = tmp_path / "novel.pdf"
    target = tmp_path / "out.txt"
    source.write_bytes(b"%PDF-1.4")

    class FakePage:
        def __init__(self, text: str):
            self.text = text

        def extract_text(self):
            return self.text

    class FakeReader:
        pages = [FakePage("第一章 风起\n你好")]

    monkeypatch.setattr(file_parser, "PdfReader", lambda _: FakeReader())

    meta = normalize_text_file(source, target)

    assert meta["parser"] == "pdf"
    assert "PDF 第 1 页" in target.read_text(encoding="utf-8")
    assert "第一章 风起" in target.read_text(encoding="utf-8")


def test_word_and_pdf_extensions_supported():
    assert validate_extension("draft.docx") == "docx"
    assert validate_extension("draft.pdf") == "pdf"


def test_legacy_doc_extension_rejected():
    with pytest.raises(HTTPException):
        validate_extension("draft.doc")
