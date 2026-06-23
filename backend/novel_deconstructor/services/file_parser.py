from __future__ import annotations

from pathlib import Path
from typing import BinaryIO

import aiofiles
from docx import Document
from fastapi import HTTPException, UploadFile
from pypdf import PdfReader


TEXT_EXTENSIONS = {".txt", ".md"}
DOCUMENT_EXTENSIONS = {".docx", ".pdf"}
SUPPORTED_EXTENSIONS = TEXT_EXTENSIONS | DOCUMENT_EXTENSIONS
ENCODINGS = ["utf-8-sig", "utf-8", "gb18030", "gbk", "big5", "utf-16", "utf-16le", "utf-16be"]


def validate_extension(filename: str) -> str:
    ext = Path(filename).suffix.lower()
    if ext == ".doc":
        raise HTTPException(status_code=400, detail="暂不支持旧版 .doc，请另存为 .docx 后上传")
    if ext not in SUPPORTED_EXTENSIONS:
        raise HTTPException(status_code=400, detail="支持 .txt、.md、.docx、.pdf 文件")
    return ext.lstrip(".")


def detect_encoding(path: Path) -> str:
    sample = path.read_bytes()[:128 * 1024]
    for encoding in ENCODINGS:
        try:
            sample.decode(encoding)
            return encoding
        except UnicodeDecodeError:
            continue
    raise ValueError("无法识别文本编码，请转为 UTF-8、GBK 或 GB18030 后重试")


def normalize_text_file(source_path: Path, output_path: Path) -> dict:
    ext = source_path.suffix.lower()
    if ext in TEXT_EXTENSIONS:
        return _normalize_plain_text(source_path, output_path)
    if ext == ".docx":
        return _extract_docx(source_path, output_path)
    if ext == ".pdf":
        return _extract_pdf(source_path, output_path)
    raise ValueError(f"不支持的文件类型: {ext}")


def _normalize_plain_text(source_path: Path, output_path: Path) -> dict:
    encoding = detect_encoding(source_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    char_count = 0
    line_count = 0

    with source_path.open("r", encoding=encoding, errors="strict", newline=None) as source:
        with output_path.open("w", encoding="utf-8", newline="\n") as target:
            for line in source:
                normalized = line.replace("\r\n", "\n").replace("\r", "\n")
                target.write(normalized)
                char_count += len(normalized)
                line_count += 1

    return {
        "encoding": encoding,
        "char_count": char_count,
        "line_count": line_count,
        "raw_path": str(output_path),
        "parser": "plain_text",
    }


def _extract_docx(source_path: Path, output_path: Path) -> dict:
    document = Document(str(source_path))
    blocks: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            blocks.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                blocks.append(" | ".join(cells))

    text = "\n\n".join(blocks).strip()
    if not text:
        raise ValueError("Word 文件未提取到可分析文本")
    return _write_extracted_text(output_path, text, "docx")


def _extract_pdf(source_path: Path, output_path: Path) -> dict:
    reader = PdfReader(str(source_path))
    pages: list[str] = []
    for page_number, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if text:
            pages.append(f"<!-- PDF 第 {page_number} 页 -->\n{text}")

    text = "\n\n".join(pages).strip()
    if not text:
        raise ValueError("PDF 未提取到可分析文本；扫描版 PDF 请先 OCR 后再上传")
    return _write_extracted_text(output_path, text, "pdf")


def _write_extracted_text(output_path: Path, text: str, parser: str) -> dict:
    normalized = text.replace("\r\n", "\n").replace("\r", "\n").strip() + "\n"
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(normalized, encoding="utf-8", newline="\n")
    return {
        "encoding": "extracted",
        "char_count": len(normalized),
        "line_count": normalized.count("\n"),
        "raw_path": str(output_path),
        "parser": parser,
    }


async def save_upload(upload: UploadFile, destination: Path, max_bytes: int) -> int:
    destination.parent.mkdir(parents=True, exist_ok=True)
    size = 0
    async with aiofiles.open(destination, "wb") as out:
        while True:
            chunk = await upload.read(1024 * 1024)
            if not chunk:
                break
            size += len(chunk)
            if size > max_bytes:
                await out.close()
                destination.unlink(missing_ok=True)
                raise HTTPException(status_code=413, detail="上传文件超过 MAX_UPLOAD_SIZE_MB 限制")
            await out.write(chunk)
    return size


def stream_copy_binary(source: BinaryIO, destination: Path, max_bytes: int) -> int:
    destination.parent.mkdir(parents=True, exist_ok=True)
    size = 0
    with destination.open("wb") as out:
        while True:
            chunk = source.read(1024 * 1024)
            if not chunk:
                break
            size += len(chunk)
            if size > max_bytes:
                destination.unlink(missing_ok=True)
                raise ValueError("文件过大")
            out.write(chunk)
    return size
